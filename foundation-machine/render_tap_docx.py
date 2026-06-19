#!/usr/bin/env python3
"""
TAP branded .docx renderer — the HARD-CODED enforcement of
reference/tap-formatting-standards.md.

Whatever the model writes, this renderer guarantees the layout rules in CODE:
  - sections never break across pages   (keepNext + keepLines per section, released
    only on the last element; cantSplit on every table row)
  - margins 0.6" top / 0.5" bottom / 0.85" left+right
  - compact 8.5-9pt fonts
  - branded cover page, gold accent bars, navy table headers with alternating
    row shading, dark callout boxes, client brand colors (form fields 48-50)

Usage (standalone test):
  python render_tap_docx.py <input.md> <output.docx> [primaryHex secondaryHex accentHex]

Called by foundation.py after a TAP passes the grade loop.
"""
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Defaults if the client gave no brand colors (Fortune 500 navy + gold).
DEFAULT_PRIMARY = "1F3A5F"   # navy
DEFAULT_SECONDARY = "2E5A88" # lighter navy
DEFAULT_ACCENT = "C5A253"    # gold
ROW_SHADE = "F2F4F7"         # light gray alternating row
CALLOUT_BG = "1F3A5F"        # dark callout background
WHITE = "FFFFFF"
BODY = "222222"


def _norm_hex(h, fallback):
    if not h:
        return fallback
    m = re.search(r"#?([0-9A-Fa-f]{6})", str(h))
    return m.group(1).upper() if m else fallback


def extract_brand_colors(answers):
    """Pull up to 3 hex colors from the submission (form fields 48-50: primary,
    secondary, accent). Falls back to navy/gold if none are present."""
    primary = secondary = accent = None
    if isinstance(answers, dict):
        # prefer explicit keys if present
        for k, v in answers.items():
            kl = k.lower()
            if primary is None and "primary" in kl and re.search(r"#?[0-9A-Fa-f]{6}", str(v)):
                primary = v
            elif secondary is None and "secondary" in kl and re.search(r"#?[0-9A-Fa-f]{6}", str(v)):
                secondary = v
            elif accent is None and "accent" in kl and re.search(r"#?[0-9A-Fa-f]{6}", str(v)):
                accent = v
        # otherwise scrape any hex codes we can find
        if not (primary and secondary and accent):
            found = []
            for v in answers.values():
                found += re.findall(r"#([0-9A-Fa-f]{6})", str(v))
            found = [f.upper() for f in found]
            if found:
                primary = primary or found[0]
                secondary = secondary or (found[1] if len(found) > 1 else found[0])
                accent = accent or (found[2] if len(found) > 2 else (found[1] if len(found) > 1 else DEFAULT_ACCENT))
    return (_norm_hex(primary, DEFAULT_PRIMARY),
            _norm_hex(secondary, DEFAULT_SECONDARY),
            _norm_hex(accent, DEFAULT_ACCENT))


# ----------------------------------------------------------------------------
# low-level docx helpers
# ----------------------------------------------------------------------------
def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement("w:cantSplit")
    trPr.append(cs)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def _set_keep(par, keep_next=True):
    """sections never break: keepLines always, keepNext unless released."""
    pf = par.paragraph_format
    pf.keep_together = True          # keepLines
    pf.keep_with_next = keep_next    # keepNext


class TapDoc:
    def __init__(self, primary, secondary, accent):
        self.doc = Document()
        self.primary, self.secondary, self.accent = primary, secondary, accent
        self._section_pars = []      # paragraphs in the current section (for keepNext release)
        self._setup()

    def _setup(self):
        for s in self.doc.sections:
            s.top_margin = Inches(0.6)
            s.bottom_margin = Inches(0.5)
            s.left_margin = Inches(0.85)
            s.right_margin = Inches(0.85)
        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(9)
        normal.font.color.rgb = RGBColor.from_string(BODY)

    # --- section bookkeeping: release keepNext on the LAST element of a section ---
    def _close_section(self):
        if self._section_pars:
            _set_keep(self._section_pars[-1], keep_next=False)
        self._section_pars = []

    def _track(self, par):
        _set_keep(par, keep_next=True)
        self._section_pars.append(par)
        return par

    # --- building blocks ---
    def gold_bar(self):
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        _no_borders(t)
        row = t.rows[0]
        _cant_split(row)
        row.height = Pt(3)
        cell = t.cell(0, 0)
        _shade(cell, self.accent)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        cell.paragraphs[0].add_run("").font.size = Pt(1)
        # keep the bar with the heading that follows
        cell.paragraphs[0].paragraph_format.keep_with_next = True

    def heading(self, text, level=2):
        self._close_section()        # new section starts here
        self.gold_bar()
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text.upper() if level <= 2 else text)
        run.bold = True
        run.font.size = Pt(13 if level <= 2 else 11)
        run.font.color.rgb = RGBColor.from_string(self.primary)
        self._track(p)

    def paragraph(self, text, bullet=False):
        p = self.doc.add_paragraph(style="List Bullet" if bullet else None)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(9)
        self._track(p)

    def callout(self, text):
        t = self.doc.add_table(rows=1, cols=1)
        _no_borders(t)
        row = t.rows[0]
        _cant_split(row)
        cell = t.cell(0, 0)
        _shade(cell, self.primary)
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(WHITE)
        # tie the callout to the surrounding section
        anchor = self.doc.add_paragraph()
        anchor.paragraph_format.space_after = Pt(0)
        self._track(anchor)

    def table(self, rows):
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        t = self.doc.add_table(rows=len(rows), cols=ncols)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.style = "Table Grid"
        for ri, rowdata in enumerate(rows):
            wrow = t.rows[ri]
            _cant_split(wrow)                       # rows never split across pages
            is_header = ri == 0
            for ci in range(ncols):
                cell = t.cell(ri, ci)
                if is_header:
                    _shade(cell, self.primary)
                elif ri % 2 == 0:
                    _shade(cell, ROW_SHADE)
                txt = rowdata[ci] if ci < len(rowdata) else ""
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_after = Pt(0)
                run = cp.add_run(txt)
                run.font.size = Pt(8.5)
                if is_header:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
        # keep the table glued to its heading: mark the paragraph before it
        if self._section_pars:
            self._section_pars[-1].paragraph_format.keep_with_next = True

    def cover(self, title, subtitle, meta_lines):
        # top brand band
        band = self.doc.add_table(rows=1, cols=1)
        _no_borders(band)
        _cant_split(band.rows[0])
        band.rows[0].height = Pt(10)
        _shade(band.cell(0, 0), self.primary)
        band.cell(0, 0).paragraphs[0].add_run("").font.size = Pt(2)

        for _ in range(4):
            self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("TARGET AUDIENCE PROFILE")
        r.bold = True
        r.font.size = Pt(30)
        r.font.color.rgb = RGBColor.from_string(self.primary)

        # gold divider
        gp = self.doc.add_paragraph()
        gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gr = gp.add_run("________________")
        gr.font.color.rgb = RGBColor.from_string(self.accent)
        gr.bold = True

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(title)
        r2.bold = True
        r2.font.size = Pt(18)
        r2.font.color.rgb = RGBColor.from_string(self.secondary)

        if subtitle:
            p3 = self.doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r3 = p3.add_run(subtitle)
            r3.font.size = Pt(12)
            r3.font.color.rgb = RGBColor.from_string(BODY)

        for _ in range(6):
            self.doc.add_paragraph()
        for line in meta_lines:
            mp = self.doc.add_paragraph()
            mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mr = mp.add_run(line)
            mr.font.size = Pt(10)
            mr.font.color.rgb = RGBColor.from_string(self.secondary)

        self.doc.add_page_break()

    def save(self, path):
        self._close_section()        # release keepNext on the final section
        self.doc.save(path)


# ----------------------------------------------------------------------------
# markdown -> TapDoc
# ----------------------------------------------------------------------------
def _clean_inline(t):
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)   # bold markers
    t = re.sub(r"`([^`]+)`", r"\1", t)
    return t.strip()


def _parse_table_block(lines):
    rows = []
    for ln in lines:
        if set(ln.replace("|", "").strip()) <= set("-: "):  # separator row
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append([_clean_inline(c) for c in cells])
    return rows


def render_markdown(md, out_path, title, subtitle="", meta_lines=None,
                    primary=DEFAULT_PRIMARY, secondary=DEFAULT_SECONDARY, accent=DEFAULT_ACCENT):
    doc = TapDoc(primary, secondary, accent)
    doc.cover(title, subtitle, meta_lines or [])

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()
        if not line:
            i += 1
            continue
        # skip the source/credit comment lines and top H1 (cover already has title)
        if line.startswith("> Source") or line.startswith("Source:") or line.startswith("Client:"):
            i += 1
            continue
        h = re.match(r"^(#{1,6})\s+(.*)", line)
        if h:
            level = len(h.group(1))
            if level == 1:   # H1 already shown on the cover
                i += 1
                continue
            doc.heading(_clean_inline(h.group(2)), level=level)
            i += 1
            continue
        if line.startswith(">"):
            doc.callout(_clean_inline(line.lstrip("> ").strip()))
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and "|" in lines[i + 1]:
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            doc.table(_parse_table_block(block))
            continue
        if re.match(r"^[-*+]\s+", line):
            doc.paragraph(_clean_inline(re.sub(r"^[-*+]\s+", "", line)), bullet=True)
            i += 1
            continue
        doc.paragraph(_clean_inline(line))
        i += 1

    doc.save(out_path)
    return out_path


def estimate_pages(md):
    """Rough page estimate for the grade-loop gate (cover + ~520 words/page compact)."""
    words = len(re.findall(r"\S+", md))
    return 1 + max(1, round(words / 520))


def _find_soffice():
    import shutil
    for c in ("soffice", "libreoffice"):
        p = shutil.which(c)
        if p:
            return p
    mac = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    import os as _os
    return mac if _os.path.exists(mac) else None


def verify_pages(docx_path):
    """TRUE page count: render the .docx to PDF with LibreOffice headless and count pages.
    Returns (pages:int, method:str) or (None, reason) when LibreOffice is not available.
    This is the authoritative 5-7 page check; the word estimate is only the in-loop proxy."""
    import os, subprocess, tempfile, shutil
    soffice = _find_soffice()
    if not soffice:
        return None, "estimate-only (install LibreOffice for exact page count)"
    try:
        tmp = tempfile.mkdtemp()
        subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", tmp, docx_path],
                       check=True, capture_output=True, timeout=120)
        pdf = os.path.join(tmp, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        if not os.path.exists(pdf):
            return None, "conversion failed"
        # count pages: macOS mdls first, then pypdf if available
        mdls = shutil.which("mdls")
        if mdls:
            out = subprocess.run([mdls, "-name", "kMDItemNumberOfPages", "-raw", pdf],
                                 capture_output=True, text=True, timeout=30).stdout.strip()
            if out.isdigit():
                return int(out), "libreoffice+mdls"
        try:
            from pypdf import PdfReader
            return len(PdfReader(pdf).pages), "libreoffice+pypdf"
        except Exception:
            return None, "rendered but no page counter (install pypdf or use macOS mdls)"
    except Exception as e:
        return None, f"verify failed: {e}"


if __name__ == "__main__":
    if len(sys.argv) < 3:
        sys.exit("usage: render_tap_docx.py <input.md> <output.docx> [primary secondary accent]")
    src, dst = sys.argv[1], sys.argv[2]
    cols = sys.argv[3:6] if len(sys.argv) >= 6 else []
    p = _norm_hex(cols[0], DEFAULT_PRIMARY) if cols else DEFAULT_PRIMARY
    s = _norm_hex(cols[1], DEFAULT_SECONDARY) if len(cols) > 1 else DEFAULT_SECONDARY
    a = _norm_hex(cols[2], DEFAULT_ACCENT) if len(cols) > 2 else DEFAULT_ACCENT
    text = open(src, encoding="utf-8").read()
    render_markdown(text, dst, title="Sample Client", subtitle="Sample Event",
                    meta_lines=["Prepared by Event Sales Agency"], primary=p, secondary=s, accent=a)
    print("wrote", dst, "| est pages:", estimate_pages(text))
